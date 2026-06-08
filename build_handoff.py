#!/usr/bin/env python3
"""Generate wellness-wheel-handoff.docx — developer brief for the Wellness Wheel prototype."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEAL = RGBColor(0x00, 0xC9, 0xB1)
NAVY = RGBColor(0x08, 0x0F, 0x1E)
COBALT = RGBColor(0x1A, 0x4F, 0xBF)
GREY = RGBColor(0x5A, 0x6A, 0x78)

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = COBALT
    return p


def para(text, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    p.add_run(text)
    return p


def kv(label, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return t


# ----------------------------- Cover -----------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Wellness Wheel")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = TEAL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Developer Handoff Brief")
r.font.size = Pt(16)
r.font.color.rgb = NAVY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Kris Pierce · Prototype to Production · Prepared for engineering")
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_paragraph()
para(
    "This document is the single source of truth for taking the Wellness Wheel "
    "prototype (wellness-wheel.jsx) from an interactive front-end demo to a "
    "production application. It covers the product, the architecture, the full "
    "gamification system, every content asset, and each backend integration point "
    "a developer needs to wire up. A developer should be able to build from this "
    "without coming back with questions.",
)
doc.add_page_break()

# ----------------------------- 1. Overview -----------------------------
h1("1. Product Overview & User Journey")
para(
    "The Wellness Wheel is a self-guided wellness assessment and habit-building "
    "experience. A user rates herself across eight dimensions of wellbeing, sees "
    "her balance visualised as a radar 'wheel', then earns points and badges by "
    "completing small activities and weekly check-ins. After ~6 weeks she "
    "reassesses and sees a before/after comparison of her growth."
)
h2("The eight dimensions")
bullet("Physical — sleep, movement, nutrition, energy")
bullet("Emotional — self-awareness, resilience, boundaries")
bullet("Intellectual — curiosity, learning, creativity")
bullet("Social — connection, belonging, relationships")
bullet("Spiritual — purpose, values, meaning, stillness")
bullet("Occupational — meaningful work, balance, growth")
bullet("Environmental — calm spaces, nature, sustainability")
bullet("Financial — control, security, progress")

h2("Core user journey")
bullet("Take the 64-question baseline assessment (8 dimensions × 8 statements, 1–5 Likert).")
bullet("View the wheel — radar chart + per-dimension scores + overall balance score.")
bullet("Complete activities from a 144-item library to earn points and level up.")
bullet("Take on optional bonus challenges (one per dimension) for larger point rewards.")
bullet("Check in weekly with a mood rating and optional note; build a streak.")
bullet("Add an accountability partner to share progress.")
bullet("After ~6 weeks, reassess and view the before/after comparison.")
bullet("Earn up to 12 badges across the whole journey.")

# ----------------------------- 2. Architecture -----------------------------
h1("2. File Structure & Component Architecture")
kv("Entry file", "wellness-wheel.jsx — a single self-contained React component, default-exported as WellnessWheel.")
kv("Dependencies", "react, recharts (radar charts), lucide-react (icons). No CSS framework required — styles are inline using the brand token object C.")
kv("Persistence", "Browser localStorage under key 'kp_wellness_wheel_v1'. This is the seam to replace with a real backend (see §5).")

h2("Component map")
table(
    ["Component", "Responsibility"],
    [
        ["WellnessWheel", "Root. Owns all state, points/level/badge derivation, persistence, routing between tabs, sticky header + nav."],
        ["Assessment", "Drives both baseline and 6-week reassessment (mode prop). 8 steps, validates all answered."],
        ["Wheel", "Radar chart, level/points card, per-dimension bars, badge shelf."],
        ["Activities", "Accordion of 8 dimensions, 18 tickable activities each."],
        ["BonusChallenges", "8 stretch challenges, one per dimension."],
        ["CheckIn", "Weekly mood (1–5 emoji) + optional note, history list, streak."],
        ["Partner", "Accountability partner email capture + validation."],
        ["Compare", "Before/after radar (two series) + per-dimension delta."],
        ["BadgeShelf", "Renders all 12 badges, earned vs locked."],
    ],
)
para(
    "State is intentionally centralised in the root component and threaded down as "
    "props. Points, level, streak, and earned badges are all derived with useMemo "
    "from primitive state (answers, completed activities/bonus, check-ins, partner) "
    "— there is no duplicated source of truth to keep in sync.",
)

# ----------------------------- 3. Gamification -----------------------------
h1("3. Gamification System")
h2("Points")
bullet("Activities: each carries 5–15 points based on effort. 144 activities total.")
bullet("Bonus challenges: 40–50 points each, 8 total.")
bullet("Weekly check-in: +5 points each.")
para("Total points = sum of completed activities + completed bonus challenges + (check-ins × 5).")

h2("Levels")
table(
    ["Level", "Name", "Points required"],
    [
        ["1", "Seedling", "0"],
        ["2", "Sprout", "100"],
        ["3", "Bloom", "300"],
        ["4", "Flourish", "600"],
        ["5", "Radiant", "1000"],
    ],
)

h2("Streaks")
para(
    "A streak counts consecutive ISO weeks with a check-in. The CheckIn component "
    "prevents more than one check-in per ISO week. Streak math lives in "
    "computeStreak/consecutive at the bottom of the file."
)

h2("Badges (12)")
table(
    ["Badge", "Unlock condition"],
    [
        ["First Steps", "Complete your first (baseline) assessment."],
        ["Wheel Turner", "Complete 10 activities."],
        ["Balanced", "Do at least one activity in every dimension."],
        ["Centurion", "Earn 100 points."],
        ["High Five", "Earn 500 points."],
        ["Level Up", "Reach Level 3 (Bloom)."],
        ["Challenger", "Complete a bonus challenge."],
        ["Overachiever", "Complete all 8 bonus challenges."],
        ["Honest With Herself", "Complete a weekly check-in. (new)"],
        ["On a Roll", "Reach a 3-week check-in streak."],
        ["Better Together", "Add an accountability partner. (new)"],
        ["Growth Visible", "Reassess and improve your overall score. (new)"],
    ],
)
para(
    "Each badge is a pure predicate over derived state (the check function in the "
    "BADGES array). Newly-earned badges trigger an unlock toast; seenBadges is "
    "persisted so a badge only celebrates once."
)

# ----------------------------- 4. Content inventory -----------------------------
h1("4. Content Inventory")
kv("Assessment questions", "64 (8 dimensions × 8 Likert statements). See QUESTIONS in the source.")
kv("Activities", "144 (18 per dimension), each with a title and point value. See ACTIVITIES.")
kv("Bonus challenges", "8 (one per dimension), 40–50 points each. See BONUS.")
kv("Badges", "12. See BADGES.")
kv("Moods", "5-point emoji scale for weekly check-ins.")
para(
    "All copy is data-driven and lives in plain JS objects at the top of the file, "
    "so content can be lifted into a CMS or database with no logic changes. Activity "
    "and question IDs are positional (e.g. 'physical:3') — preserve ordering, or "
    "migrate to stable IDs before launch."
)

# ----------------------------- 5. Backend integration -----------------------------
h1("5. Backend Integration Points")
para("The prototype is fully client-side. These are the seams to productionise, in priority order.", bold=True)

h2("5.1 Authentication")
bullet("Add user accounts (email/password or social). Gate the whole experience behind auth.")
bullet("Replace the anonymous localStorage profile with a per-user server profile.")

h2("5.2 Database / persistence")
para("Replace the localStorage read/write (loadState + the persistence useEffect) with API calls. Suggested schema:")
table(
    ["Table", "Key columns"],
    [
        ["users", "id, email, name, created_at, partner_email"],
        ["assessments", "id, user_id, type ('baseline'|'reassess'), answers (jsonb), created_at"],
        ["activity_log", "id, user_id, activity_key, points, completed_at"],
        ["bonus_log", "id, user_id, bonus_key, points, completed_at"],
        ["checkins", "id, user_id, iso_week, mood (1–5), note, created_at"],
        ["badges", "id, user_id, badge_id, earned_at"],
    ],
)
para("Points, level, streak, and badges should be derived server-side from these tables to prevent client tampering.")

h2("5.3 Accountability partner email")
bullet("On partner save, send an invite/confirmation email.")
bullet("Send a weekly digest (mood, points gained, badges earned) to the partner — cron/scheduled job.")
bullet("Honour unsubscribe and store partner consent.")

h2("5.4 Push notifications & reminders")
bullet("Weekly check-in reminder (respect the user's timezone and ISO-week boundary).")
bullet("Streak-at-risk nudge.")
bullet("6-week reassessment prompt, timed from the baseline date.")

h2("5.5 Analytics events")
para("Instrument at minimum: assessment_started, assessment_completed, activity_completed, bonus_completed, checkin_saved, partner_added, badge_earned, reassessment_completed, level_up.")

h2("5.6 Monetisation gate (optional)")
para(
    "If a paid tier is desired, a natural gate is: baseline assessment + the wheel "
    "are free; the full activity library, bonus challenges, partner, and reassessment "
    "comparison sit behind a membership. The KrisPierce site already has a membership "
    "page to align with."
)

# ----------------------------- 6. Local dev -----------------------------
h1("6. Local Development Setup")
h2("Option A — Vite (recommended)")
bullet("npm create vite@latest wellness -- --template react")
bullet("npm i recharts lucide-react")
bullet("Drop wellness-wheel.jsx into src/, import and render <WellnessWheel /> in App.jsx.")
bullet("npm run dev")
h2("Option B — Create React App")
bullet("npx create-react-app wellness && cd wellness")
bullet("npm i recharts lucide-react, then import the component.")
h2("Option C — claude.ai artifact")
bullet("The component is artifact-compatible: it uses only react, recharts, and lucide-react and has no external CSS. Paste it directly.")

# ----------------------------- 7. Roadmap -----------------------------
h1("7. Recommended Build Roadmap")
h2("Phase 1 — Foundation")
bullet("Auth + user profiles.")
bullet("Move assessment + wheel to server-backed persistence.")
bullet("Server-side score/level derivation.")
h2("Phase 2 — Engagement")
bullet("Activities, bonus challenges, points, badges persisted and validated server-side.")
bullet("Weekly check-ins + streaks + reminder notifications.")
bullet("Analytics instrumentation.")
h2("Phase 3 — Social & growth")
bullet("Accountability partner emails + weekly digest.")
bullet("6-week reassessment flow + before/after comparison.")
bullet("Optional membership gate + monetisation.")

# ----------------------------- 8. Limitations -----------------------------
h1("8. Known Limitations & Design Decisions")
bullet("All state is local to the browser; clearing storage resets everything. By design for a prototype.")
bullet("Activity/question identity is positional — adopt stable IDs before a CMS migration.")
bullet("No anti-cheat: a user could mark activities without doing them. Acceptable for self-guided wellness; tighten only if points become monetised/competitive.")
bullet("Reassessment overwrites 'current' each time; only baseline vs latest is compared. Add an assessment history table for trend lines if desired.")
bullet("Partner email is captured but not yet sent anything — the integration is stubbed with a note in the UI.")
bullet("Streak uses ISO weeks; a user in a different timezone near a week boundary may see an off-by-one. Derive server-side with the user's timezone in production.")

# ----------------------------- 9. Design specs -----------------------------
h1("9. Design Tokens & Specs")
h2("Brand colour palette")
table(
    ["Token", "Hex", "Use"],
    [
        ["Vivid teal", "#00C9B1", "Primary accent, progress, earned states"],
        ["Aqua", "#3DD6C8", "Secondary accent"],
        ["Seafoam", "#80EAD8", "Soft highlights, success text"],
        ["Pale mint", "#C2F5EE", "Badge / toast text on dark"],
        ["Ghost", "#EEF8F7", "Primary text on dark"],
        ["Cobalt", "#1A4FBF", "Gradient partner to teal"],
        ["Ocean", "#0E7FA8", "Tertiary accent"],
        ["Deep indigo", "#2D3580", "Depth accent"],
        ["Midnight navy", "#080F1E", "Page background"],
        ["Navy card", "#0E1C35", "Card surface"],
        ["Navy lift", "#152844", "Inputs, track backgrounds"],
        ["Hairline", "#1D2F4D", "Borders, dividers"],
        ["Warm grey", "#8BA8A5", "Muted / secondary text"],
    ],
)
h2("Layout & type")
bullet("Max content width 1100px, centred.")
bullet("Cards: 16px radius, 1px hairline border, 20px padding.")
bullet("Pills/buttons: 999px radius. Primary = teal→cobalt gradient on navy text.")
bullet("Font: Inter / system-ui. Headings 800 weight.")
bullet("Each dimension has its own accent colour for quick visual scanning.")
h2("Accessibility notes")
bullet("Ensure 1–5 rating buttons and mood emojis have aria-labels in production.")
bullet("Verify contrast of grey (#8BA8A5) on navy for body text; bump to ghost where it carries meaning.")
bullet("Make the radar chart data available as a table for screen readers (the per-dimension bars already serve this).")
bullet("Support keyboard navigation and visible focus rings on all interactive pills.")

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("End of brief — Wellness Wheel handoff")
r.italic = True
r.font.color.rgb = GREY

out = "/home/user/KrisPierce/wellness-wheel-handoff.docx"
doc.save(out)
print("Saved", out)
